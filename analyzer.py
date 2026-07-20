import re
from collections import Counter
from math import sqrt
from pathlib import Path

from docx import Document
from pypdf import PdfReader


TECHNICAL_SKILLS = {
    "python",
    "java",
    "javascript",
    "typescript",
    "html",
    "css",
    "react",
    "angular",
    "vue",
    "node.js",
    "express",
    "flask",
    "django",
    "fastapi",
    "sql",
    "mysql",
    "postgresql",
    "mongodb",
    "sqlite",
    "git",
    "github",
    "docker",
    "kubernetes",
    "aws",
    "azure",
    "gcp",
    "rest api",
    "graphql",
    "machine learning",
    "deep learning",
    "data science",
    "data analysis",
    "pandas",
    "numpy",
    "tensorflow",
    "pytorch",
    "scikit-learn",
    "power bi",
    "tableau",
    "figma",
    "ui/ux",
    "cybersecurity",
    "linux",
    "cloud computing",
    "api integration",
    "responsive design",
    "software engineering",
    "object oriented programming",
    "oop",
    "firebase",
    "bootstrap",
    "tailwind css",
    "jenkins",
    "ci/cd",
    "devops",
    "testing",
    "unit testing",
}

PROFESSIONAL_SKILLS = {
    "communication",
    "leadership",
    "problem solving",
    "teamwork",
    "time management",
    "project management",
    "critical thinking",
    "collaboration",
    "agile",
    "scrum",
    "adaptability",
    "creativity",
    "research",
    "presentation",
}

ALL_SKILLS = TECHNICAL_SKILLS | PROFESSIONAL_SKILLS

ACTION_VERBS = {
    "achieved",
    "analyzed",
    "automated",
    "built",
    "collaborated",
    "created",
    "delivered",
    "deployed",
    "designed",
    "developed",
    "engineered",
    "implemented",
    "improved",
    "increased",
    "integrated",
    "launched",
    "led",
    "maintained",
    "managed",
    "optimized",
    "organized",
    "reduced",
    "resolved",
    "tested",
}

SECTION_PATTERNS = {
    "contact information": (
        r"\b(email|phone|linkedin|github)\b"
    ),
    "professional summary": (
        r"\b(summary|profile|objective|about me)\b"
    ),
    "experience": (
        r"\b(experience|employment|work history|internship)\b"
    ),
    "education": (
        r"\b(education|academic|university|college|degree)\b"
    ),
    "skills": (
        r"\b(skills|technologies|technical skills|competencies)\b"
    ),
    "projects": (
        r"\b(projects|portfolio|case studies)\b"
    ),
}

STOP_WORDS = {
    "and",
    "the",
    "with",
    "for",
    "from",
    "that",
    "this",
    "your",
    "you",
    "are",
    "our",
    "will",
    "have",
    "has",
    "into",
    "using",
    "work",
    "role",
    "job",
    "team",
    "skills",
    "experience",
    "years",
    "about",
    "who",
    "all",
    "but",
    "not",
    "can",
    "their",
    "they",
    "was",
    "were",
    "been",
    "being",
    "than",
    "then",
    "also",
    "should",
    "must",
    "required",
    "preferred",
    "responsibilities",
    "requirements",
}


def extract_resume_text(
    file_path: Path,
    extension: str,
) -> str:
    extension = extension.lower().strip(".")

    if extension == "pdf":
        return extract_pdf_text(file_path)

    if extension == "docx":
        return extract_docx_text(file_path)

    if extension == "txt":
        return file_path.read_text(
            encoding="utf-8",
            errors="ignore",
        )

    raise ValueError(
        "Unsupported resume format. Please upload PDF, DOCX or TXT."
    )


def extract_pdf_text(file_path: Path) -> str:
    try:
        reader = PdfReader(str(file_path))

        pages = [
            page.extract_text() or ""
            for page in reader.pages
        ]

        return "\n".join(pages)

    except Exception as error:
        raise ValueError(
            "The PDF could not be read. It may be damaged, encrypted or image-only."
        ) from error


def extract_docx_text(file_path: Path) -> str:
    try:
        document = Document(str(file_path))

        paragraphs = [
            paragraph.text
            for paragraph in document.paragraphs
        ]

        table_text = []

        for table in document.tables:
            for row in table.rows:
                table_text.extend(
                    cell.text
                    for cell in row.cells
                )

        return "\n".join(
            paragraphs + table_text
        )

    except Exception as error:
        raise ValueError(
            "The DOCX file could not be read."
        ) from error


def normalize_text(text: str) -> str:
    return re.sub(
        r"\s+",
        " ",
        text.lower(),
    ).strip()


def tokenize(text: str) -> list[str]:
    return re.findall(
        r"[a-zA-Z][a-zA-Z0-9+.#/-]{1,}",
        text.lower(),
    )


def detect_skills(text: str) -> list[str]:
    normalized_text = normalize_text(text)
    detected = []

    for skill in ALL_SKILLS:
        pattern = (
            rf"(?<!\w)"
            rf"{re.escape(skill)}"
            rf"(?!\w)"
        )

        if re.search(
            pattern,
            normalized_text,
        ):
            detected.append(skill)

    return sorted(detected)


def important_keywords(text: str) -> list[str]:
    words = tokenize(text)

    filtered_words = [
        word
        for word in words
        if word not in STOP_WORDS
        and len(word) > 2
    ]

    frequency = Counter(filtered_words)

    return [
        word
        for word, _count
        in frequency.most_common(45)
    ]


def semantic_similarity(
    resume_text: str,
    job_description: str,
) -> float:
    if not job_description.strip():
        return 0.0

    resume_words = [
        word
        for word in tokenize(resume_text)
        if word not in STOP_WORDS
    ]

    job_words = [
        word
        for word in tokenize(job_description)
        if word not in STOP_WORDS
    ]

    if not resume_words or not job_words:
        return 0.0

    resume_counts = Counter(resume_words)
    job_counts = Counter(job_words)

    shared_words = (
        set(resume_counts)
        & set(job_counts)
    )

    dot_product = sum(
        resume_counts[word]
        * job_counts[word]
        for word in shared_words
    )

    resume_length = sqrt(
        sum(
            count ** 2
            for count in resume_counts.values()
        )
    )

    job_length = sqrt(
        sum(
            count ** 2
            for count in job_counts.values()
        )
    )

    if resume_length == 0 or job_length == 0:
        return 0.0

    similarity = dot_product / (
        resume_length * job_length
    )

    return max(
        0.0,
        min(1.0, similarity),
    )


def calculate_job_match(
    resume_text: str,
    job_description: str,
) -> tuple[int, list[str], list[str]]:
    if not job_description.strip():
        return 0, [], []

    resume_keywords = set(
        important_keywords(resume_text)
    )

    job_keywords = set(
        important_keywords(job_description)
    )

    matching_keywords = sorted(
        resume_keywords & job_keywords
    )

    keyword_score = (
        len(matching_keywords)
        / max(len(job_keywords), 1)
    ) * 100

    resume_skills = set(
        detect_skills(resume_text)
    )

    job_skills = set(
        detect_skills(job_description)
    )

    matching_job_skills = (
        resume_skills & job_skills
    )

    missing_skills = sorted(
        job_skills - resume_skills
    )

    if job_skills:
        skill_score = (
            len(matching_job_skills)
            / len(job_skills)
        ) * 100
    else:
        skill_score = keyword_score

    semantic_score = (
        semantic_similarity(
            resume_text,
            job_description,
        )
        * 100
    )

    final_match = round(
        semantic_score * 0.50
        + keyword_score * 0.30
        + skill_score * 0.20
    )

    final_match = min(
        100,
        max(0, final_match),
    )

    return (
        final_match,
        matching_keywords,
        missing_skills,
    )


def analyze_resume(
    resume_text: str,
    job_description: str = "",
) -> dict:
    normalized_resume = normalize_text(
        resume_text
    )

    resume_words = tokenize(resume_text)
    word_count = len(resume_words)

    detected_skills = detect_skills(
        resume_text
    )

    detected_sections = [
        section
        for section, pattern
        in SECTION_PATTERNS.items()
        if re.search(
            pattern,
            normalized_resume,
        )
    ]

    action_verbs = sorted(
        {
            verb
            for verb in ACTION_VERBS
            if verb in resume_words
        }
    )

    measurable_results = re.findall(
        r"\b\d+(?:\.\d+)?%"
        r"|\b\d+\+?\b",
        resume_text,
    )

    email_found = bool(
        re.search(
            r"[\w.+-]+@[\w-]+\.[\w.-]+",
            resume_text,
        )
    )

    phone_found = bool(
        re.search(
            r"(?:\+?\d[\d\s()-]{7,}\d)",
            resume_text,
        )
    )

    (
        job_match,
        matching_keywords,
        missing_skills,
    ) = calculate_job_match(
        resume_text,
        job_description,
    )

    if not job_description.strip():
        recommended_skills = {
            "communication",
            "git",
            "problem solving",
            "sql",
            "teamwork",
        }

        missing_skills = sorted(
            recommended_skills
            - set(detected_skills)
        )

    section_score = round(
        (
            len(detected_sections)
            / len(SECTION_PATTERNS)
        )
        * 22
    )

    skill_score = min(
        20,
        len(detected_skills) * 2,
    )

    if 350 <= word_count <= 900:
        length_score = 10

    elif 220 <= word_count <= 1100:
        length_score = 7

    else:
        length_score = 3

    impact_score = min(
        15,
        len(action_verbs)
        + min(
            len(measurable_results),
            7,
        ),
    )

    contact_score = 0

    if email_found:
        contact_score += 5

    if phone_found:
        contact_score += 5

    if job_description.strip():
        relevance_score = round(
            job_match * 0.23
        )
    else:
        relevance_score = 16

    final_score = (
        section_score
        + skill_score
        + length_score
        + impact_score
        + contact_score
        + relevance_score
    )

    final_score = max(
        0,
        min(100, final_score),
    )

    strengths = []

    if len(detected_sections) >= 5:
        strengths.append(
            "Your resume includes most essential recruiter-friendly sections."
        )

    if len(detected_skills) >= 7:
        strengths.append(
            "You present a useful mix of technical and professional skills."
        )

    if len(action_verbs) >= 5:
        strengths.append(
            "Your content uses strong action-oriented language."
        )

    if measurable_results:
        strengths.append(
            "You include measurable figures that make achievements more credible."
        )

    if email_found and phone_found:
        strengths.append(
            "Your essential contact information is clearly detectable."
        )

    if (
        job_description.strip()
        and job_match >= 55
    ):
        strengths.append(
            "Your resume has a solid match with the selected job description."
        )

    if not strengths:
        strengths.append(
            "Your resume provides a useful foundation that can be improved with clearer evidence."
        )

    improvements = []

    missing_sections = [
        section
        for section in SECTION_PATTERNS
        if section not in detected_sections
    ]

    if missing_sections:
        improvements.append(
            "Add or clearly label these sections: "
            + ", ".join(missing_sections[:3])
            + "."
        )

    if len(action_verbs) < 4:
        improvements.append(
            "Start more bullet points with strong verbs such as developed, led, improved or automated."
        )

    if not measurable_results:
        improvements.append(
            "Add numbers, percentages, users, project scale or time saved to prove impact."
        )

    if word_count < 300:
        improvements.append(
            "Add more evidence about projects, responsibilities and outcomes."
        )

    elif word_count > 1000:
        improvements.append(
            "Reduce repetition and keep the resume concise, ideally one or two pages."
        )

    if not email_found or not phone_found:
        improvements.append(
            "Ensure your email address and phone number are clearly visible."
        )

    if missing_skills:
        improvements.append(
            "Consider demonstrating relevant experience with: "
            + ", ".join(missing_skills[:5])
            + "."
        )

    if (
        job_description.strip()
        and job_match < 45
    ):
        improvements.append(
            "Tailor your summary, skills and experience bullets more closely to the job description."
        )

    return {
        "score": final_score,
        "job_match": job_match,
        "word_count": word_count,
        "keyword_count": len(
            matching_keywords
        ),
        "detected_skills": detected_skills,
        "missing_skills": missing_skills,
        "strengths": strengths[:6],
        "improvements": improvements[:7],
    }